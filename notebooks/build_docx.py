# Build Word (.docx) versions of both manuscripts (bioRxiv-accepted editable format).
# Lightweight markdown -> docx: headings, bold/italic/code runs, tables, figures.
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PAPERS = {
    "paper1": {
        "md": "manuscript/manuscript_draft.md",
        "out": "manuscript/CSC_marker_benchmark.docx",
        "figdir": "manuscript/figures",
        "figures": [("Figure1.png","Figure 1.","Single-cell identification of breast cancer stem cells (GSE176078)."),
                    ("Figure2.png","Figure 2.","Functional benchmarking of CSC marker-ranking methods."),
                    ("Figure3.png","Figure 3.","Multi-cancer functional benchmark: marker rankings are tissue-specific."),
                    ("Figure4.png","Figure 4.","Cross-cancer validation identifies a recurrent CSC program."),
                    ("Figure5.png","Figure 5.","Subtype-stratified survival association of CSC signatures (METABRIC).")],
    },
    "paper2": {
        "md": "manuscript2/immune_manuscript_draft.md",
        "out": "manuscript2/CSC_immune.docx",
        "figdir": "manuscript2/figures",
        "figures": [("Figure1.png","Figure 1.","TNBC cancer stem cells upregulate CD47 and retain antigen presentation "
                     "(subtype-specific, independently replicated).")],
    },
}

def add_runs(par, text):
    """Add text with **bold**, *italic*, `code` inline formatting."""
    for tok in re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text):
        if not tok: continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = par.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Courier New"
        else:
            par.add_run(tok)

def build(cfg):
    text = open(cfg["md"]).read()
    text = re.sub(r"^<!--.*?-->\s*", "", text, count=1, flags=re.DOTALL)
    lines = text.split("\n")

    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Times New Roman"; style.font.size = Pt(11)

    in_titleblock = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # tables: consecutive lines starting with |
        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].rstrip().startswith("|"):
                block.append(lines[i].rstrip()); i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in block
                    if not re.match(r"^\|[\s:|-]+\|?$", r)]
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(t.rows[ri].cells):
                            p = t.rows[ri].cells[ci].paragraphs[0]
                            add_runs(p, cell)
                            if ri == 0:
                                for rn in p.runs: rn.bold = True
            continue
        if line.startswith("# "):
            h = doc.add_heading(line[2:], level=0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            in_titleblock = True
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1); in_titleblock = False
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.strip() == "---":
            in_titleblock = False
        elif line.strip() == "":
            pass
        else:
            p = doc.add_paragraph()
            if in_titleblock:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(p, line)
        i += 1

    # figures
    doc.add_page_break()
    doc.add_heading("Figures", level=1)
    for fn, num, cap in cfg["figures"]:
        p = os.path.join(cfg["figdir"], fn)
        if os.path.exists(p):
            doc.add_picture(p, width=Inches(6.3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp = doc.add_paragraph(); r = cp.add_run(num + " "); r.bold = True
            cp.add_run(cap); cp.paragraph_format.space_after = Pt(10)

    doc.save(cfg["out"])
    print(f"Saved {cfg['out']}  ({os.path.getsize(cfg['out'])/1024:.0f} KB)")

for cfg in PAPERS.values():
    build(cfg)
